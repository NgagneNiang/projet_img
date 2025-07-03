# # Version corrigée du script catalogue.py
# import os
# from PIL import Image
# import docx
# from docx.shared import Inches
# import pandas as pd
# from docx.oxml import parse_xml
# from docx.oxml.ns import nsdecls
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# # === Paramètres ===
# root_path = r"F:\projet_img"

# # === Extraction des images avec structure à 3 niveaux ===
# def process_strate(strate_path):
#     images = []
#     print(f"  📁 Traitement du dossier: {strate_path}")
    
#     # Vérifier si le dossier existe
#     if not os.path.exists(strate_path):
#         print(f"  ❌ Dossier inexistant: {strate_path}")
#         return images
    
#     # Parcourir les dossiers de produits
#     for product_dir in os.listdir(strate_path):
#         product_path = os.path.join(strate_path, product_dir)
#         if not os.path.isdir(product_path):
#             continue
            
#         print(f"    📂 Produit: {product_dir}")
        
#         # Parcourir les dossiers d'unités
#         for unit_dir in os.listdir(product_path):
#             unit_path = os.path.join(product_path, unit_dir)
#             if not os.path.isdir(unit_path):
#                 continue
                
#             print(f"      📦 Unité: {unit_dir}")
            
#             # Rechercher directement les images dans le dossier unité
#             best_img = None
#             best_size = 0
#             images_found = 0
            
#             for file in os.listdir(unit_path):
#                 if file.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff')):
#                     img_path = os.path.join(unit_path, file)
#                     try:
#                         file_size = os.path.getsize(img_path)
#                         images_found += 1
#                         if file_size > best_size:
#                             best_size = file_size
#                             best_img = img_path
#                         print(f"        🖼️ Image trouvée: {file} ({file_size} bytes)")
#                     except Exception as e:
#                         print(f"        ❌ Erreur avec {file}: {e}")
#                         continue

#             if best_img:
#                 images.append({
#                     'path': best_img,
#                     'libelle_groupe': product_dir.strip(),
#                     'libelle_pdt': product_dir.strip(),  # Nom du produit
#                     'libelle_unite': unit_dir.strip()    # Unité
#                 })
#                 print(f"        ✅ Meilleure image sélectionnée: {os.path.basename(best_img)}")
#             else:
#                 print(f"        ⚠️ Aucune image trouvée dans {unit_dir}")
    
#     print(f"  📊 Total images trouvées pour cette strate: {len(images)}")
#     return images

# # === Fonction pour grouper les images par produit ===
# def group_images_by_product(images_list):
#     """Groupe les images par produit pour créer des pages séparées"""
#     grouped = {}
#     for img in images_list:
#         product = img['libelle_groupe']
#         if product not in grouped:
#             grouped[product] = []
#         grouped[product].append(img)
#     return grouped

# # === Redimensionnement d'image ===
# def resize_image_to_fixed_size(img_path, target_width_cm=7, target_height_cm=5):
#     """Redimensionne l'image en gardant les proportions"""
#     try:
#         with Image.open(img_path) as img:
#             # Convertir cm en pixels (approximation 96 DPI)
#             target_width_px = int(target_width_cm * 37.8)  # 96 DPI
#             target_height_px = int(target_height_cm * 37.8)
            
#             # Calculer les nouvelles dimensions en gardant les proportions
#             img_ratio = img.width / img.height
#             target_ratio = target_width_px / target_height_px
            
#             if img_ratio > target_ratio:
#                 # Image plus large - ajuster par la largeur
#                 new_width = target_width_px
#                 new_height = int(target_width_px / img_ratio)
#             else:
#                 # Image plus haute - ajuster par la hauteur
#                 new_height = target_height_px
#                 new_width = int(target_height_px * img_ratio)
            
#             # Redimensionner avec une bonne qualité
#             img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
#             # Sauvegarder temporairement (ou créer une copie)
#             temp_path = img_path.replace('.jpg', '_resized.jpg').replace('.png', '_resized.png')
#             img_resized.save(temp_path, quality=95)
#             return temp_path
#     except Exception as e:
#         print(f"    ❌ Erreur redimensionnement {img_path}: {e}")
#         return img_path  # Retourner l'original en cas d'erreur

# # === Créer un tableau sans bordures visibles ===
# def set_no_borders(table):
#     """Supprime les bordures du tableau"""
#     tbl = table._tbl
#     for row in tbl.tr_lst:
#         for cell in row.tc_lst:
#             tcPr = cell.tcPr
#             tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
#             tcPr.append(tcBorders)

# # === Création du catalogue Word ===
# def create_word_catalog(images_list, strate_name, output_filename):
#     """Crée le catalogue Word avec les images groupées par produit"""
#     if not images_list:
#         print(f"  ⚠️ Aucune image à traiter pour {strate_name}")
#         return
    
#     doc = docx.Document()

#     # Marges étroites
#     sections = doc.sections
#     for section in sections:
#         section.left_margin = Inches(0.5)
#         section.right_margin = Inches(0.5)
#         section.top_margin = Inches(0.5)
#         section.bottom_margin = Inches(0.5)

#     # Grouper les images par produit
#     grouped_images = group_images_by_product(images_list)
    
#     print(f"  📑 Création du catalogue avec {len(grouped_images)} groupes de produits")
    
#     # 6 images par page (3 lignes x 2 colonnes)
#     rows_per_page = 3
#     cols_per_page = 2
#     images_per_page = rows_per_page * cols_per_page

#     page_count = 0
    
#     # Pour chaque groupe de produits
#     for product_name, product_images in grouped_images.items():
#         print(f"    📄 Traitement du produit: {product_name} ({len(product_images)} images)")
        
#         # Traiter les images par pages de 6
#         for i in range(0, len(product_images), images_per_page):
#             page_images = product_images[i:i + images_per_page]
#             page_count += 1
            
#             # Titre de la page
#             title_paragraph = doc.add_paragraph()
#             title_run = title_paragraph.add_run(f'{strate_name} - {product_name}')
#             title_run.font.size = docx.shared.Pt(16)
#             title_run.bold = True
#             title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
#             # Créer le tableau
#             table = doc.add_table(rows=rows_per_page, cols=cols_per_page)
#             table.style = 'Table Grid'
#             set_no_borders(table)
            
#             # Remplir le tableau
#             for row in range(rows_per_page):
#                 for col in range(cols_per_page):
#                     idx = row * cols_per_page + col
#                     if idx < len(page_images):
#                         img_info = page_images[idx]
#                         cell = table.cell(row, col)
                        
#                         # Nettoyer la cellule
#                         cell.text = ''
                        
#                         # Ajouter le titre
#                         title_para = cell.add_paragraph()
#                         title_para.text = f'{img_info["libelle_pdt"]} - {img_info["libelle_unite"]}'
#                         title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
#                         # Redimensionner et ajouter l'image
#                         try:
#                             resized_img_path = resize_image_to_fixed_size(img_info['path'])
#                             img_para = cell.add_paragraph()
#                             img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                             img_run = img_para.add_run()
#                             img_run.add_picture(resized_img_path, width=Inches(2.5), height=Inches(1.8))
#                             print(f"      ✅ Image ajoutée: {os.path.basename(img_info['path'])}")
#                         except Exception as e:
#                             print(f"      ❌ Erreur ajout image {img_info['path']}: {e}")
#                             # Ajouter un texte de remplacement
#                             error_para = cell.add_paragraph()
#                             error_para.text = "[Image non disponible]"
#                             error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
#             # Saut de page sauf pour la dernière page
#             if i + images_per_page < len(product_images) or product_name != list(grouped_images.keys())[-1]:
#                 doc.add_page_break()
    
#     # Sauvegarder le document
#     try:
#         doc.save(output_filename)
#         print(f"  ✅ Catalogue sauvegardé: {output_filename} ({page_count} pages)")
#     except Exception as e:
#         print(f"  ❌ Erreur sauvegarde: {e}")

# # === Liste des strates ===
# strates = [
#     "101_KOLDA"
#     # "102_VELINGARA",
#     # "103_MEDINA YORO FOULAH",
#     # "111_MATAM",
#     # "112_KANEL",
#     # "113_RANEROU",
#     # "11_DAKAR",
#     # "121_KAFFRINE",
#     # "122_BIRKELANE",
#     # "123_KOUNGHEUL",
#     # "124_MALEM HODDAR",
#     # "12_PIKINE",
#     # "131_KEDOUGOU",
#     # "132_SALEMATA",
#     # "133_SARAYA",
#     # "13_RUFISQUE",
#     # "141_SEDHIOU",
#     # "142_BOUNKILING",
#     # "143_GOUDOMP",
#     # "14_GUEDIAWAYE",
#     # "15_KEUR MASSAR",
#     # "21_BIGNONA",
#     # "22_OUSSOUYE",
#     # "23_ZIGUINCHOR",
#     # "31_BAMBEY",
#     # "32_DIOURBEL",
#     # "33_MBACKE",
#     # "41_DAGANA",
#     # "42_PODOR",
#     # "43_SAINT-LOUIS",
#     # "51_BAKEL",
#     # "52_TAMBACOUNDA",
#     # "53_GOUDIRY",
#     # "54_KOUMPENTOUM",
#     # "61_KAOLACK",
#     # "62_NIORO",
#     # "63_GUINGUINEO",
#     # "71_MBOUR",
#     # "72_THIES",
#     # "73_TIVAOUANE",
#     # "81_KEBEMER",
#     # "82_LINGUERE",
#     # "83_LOUGA",
#     # "91_FATICK",
#     # "92_FOUNDIOUGNE",
#     # "93_GOSSAS",
# ]

# # === Lancement ===
# print("🚀 Démarrage de la génération des catalogues...")
# print(f"📂 Répertoire racine: {root_path}")

# total_processed = 0
# total_errors = 0

# for strate in strates:
#     strate_path = os.path.join(root_path, strate)
#     if os.path.exists(strate_path):
#         print(f"\n▶️ Traitement de : {strate}")
#         try:
#             images = process_strate(strate_path)
            
#             if images:
#                 # Déterminer le chemin de sortie
#                 output_file = os.path.join(strate_path, f"{strate.replace(' ', '_')}_catalogue.docx")
                
#                 # Créer et enregistrer le fichier Word
#                 create_word_catalog(images, strate, output_file)
#                 total_processed += 1
#             else:
#                 print(f"  ⚠️ Aucune image trouvée dans {strate}")
                
#         except Exception as e:
#             print(f"  ❌ Erreur lors du traitement de {strate}: {e}")
#             total_errors += 1
#     else:
#         print(f"❌ Dossier introuvable : {strate_path}")
#         total_errors += 1

# print(f"\n✅ Génération terminée!")
# print(f"📊 Statistiques:")
# print(f"   - Strates traitées avec succès: {total_processed}")
# print(f"   - Erreurs rencontrées: {total_errors}")
# print(f"   - Total strates: {len(strates)}")
# catalogue1.py (version simplifiée pour utiliser les images préparées)
import os
import glob
from PIL import Image, UnidentifiedImageError
import docx
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from math import ceil
import traceback
import shutil
import time

# === Paramètres ===
root_path = r"F:\projet\projet_img"
DPI = 96
LOG_FILE = "log_erreurs.txt"
MAX_IMAGE_SIZE_MB = 10  # Taille maximale d'image en MB

# Dossier temporaire unique pour stocker les images redimensionnées
TEMP_DIR = os.path.join(root_path, "temp_image_processing")

def process_strate(strate_path):
    images = []
    if not os.path.exists(strate_path):
        return images
    
    print(f"\n▶️ Traitement de la strate : {strate_path}")
    total_images = 0
    valid_images = 0
    
    for product_dir in sorted(os.listdir(strate_path)):
        product_path = os.path.join(strate_path, product_dir)
        if not os.path.isdir(product_path):
            continue
            
        product_name = product_dir.strip()
        print(f"\n📁 Produit: {product_name}")
        
        for unit_dir in sorted(os.listdir(product_path)):
            unit_path = os.path.join(product_path, unit_dir)
            if not os.path.isdir(unit_path):
                continue
                
            unit_name = unit_dir.strip()
            print(f"  🖼 Unité: {unit_name}")
            
            potential_images = []
            for file in os.listdir(unit_path):
                if file.lower().endswith(('.jpg', '.jpeg', '.png')):
                    img_path = os.path.join(unit_path, file)
                    try:
                        file_size = os.path.getsize(img_path)
                        # Vérification de la taille du fichier
                        if file_size > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                            print(f"    ⚠️ Image trop grande ({file_size/1024/1024:.2f} MB > {MAX_IMAGE_SIZE_MB} MB): {file}")
                            log_error(img_path, f"Image trop grande ({file_size/1024/1024:.2f} MB)")
                            continue
                        potential_images.append((img_path, file_size))
                        total_images += 1
                    except OSError as e:
                        log_error(img_path, f"Erreur accès fichier: {str(e)}")
                        continue
            
            potential_images.sort(key=lambda x: x[1], reverse=True)
            found_valid_image = False
            
            for img_path, size in potential_images:
                try:
                    start_time = time.time()
                    with Image.open(img_path) as img_test:
                        img_test.verify()
                    
                    # Vérification supplémentaire de l'image
                    img_test = Image.open(img_path)
                    img_test.load()
                    
                    images.append({
                        'path': img_path, 
                        'libelle_produit': product_name, 
                        'libelle_unite': unit_name,
                        'size': f"{img_test.size[0]}x{img_test.size[1]}",
                        'format': img_test.format
                    })
                    valid_images += 1
                    found_valid_image = True
                    print(f"    ✅ Image valide: {os.path.basename(img_path)} ({img_test.size[0]}x{img_test.size[1]}, {img_test.format})")
                    break
                    
                except Exception as e:
                    error_type = type(e).__name__
                    print(f"    ❌ Erreur image {os.path.basename(img_path)}: {error_type} - {str(e)}")
                    log_error(img_path, f"{error_type}: {str(e)}")
                    continue
            
            if not found_valid_image and potential_images:
                log_error(unit_path, "Aucune image valide n'a pu être lue")
                print(f"    ⚠️ Aucune image valide trouvée pour {unit_name}")

    print(f"\n📊 Résumé pour {strate_path}:")
    print(f"- Images trouvées: {total_images}")
    print(f"- Images valides: {valid_images}")
    print(f"- Taux de succès: {valid_images/total_images*100:.1f}%" if total_images > 0 else "Aucune image trouvée")
    
    return images

def group_images_by_product(images_list):
    grouped = {}
    for img in images_list:
        product = img['libelle_produit']
        if product not in grouped:
            grouped[product] = []
        grouped[product].append(img)
    return grouped

def log_error(path, message):
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {path} => {message}\n")

def set_table_borders_invisible(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'nil')
        tblBorders.append(border_el)

def process_and_resize_image(img_path, target_width_inches, unique_id):
    """
    Redimensionne une image et la sauvegarde dans le dossier temporaire
    avec un nom unique. Gère la transparence PNG en fond blanc.
    """
    try:
        print(f"\n🔄 Traitement de l'image #{unique_id}: {os.path.basename(img_path)}")
        temp_filename = f"temp_image_{unique_id}.jpeg"
        temp_path = os.path.join(TEMP_DIR, temp_filename)

        with Image.open(img_path) as img:
            print(f"  Original: {img.size}, {img.format}, {img.mode}")
            target_width_px = int(target_width_inches * DPI)
            print(f"  Cible: {target_width_px}px de large")
            
            # Calcul du ratio de redimensionnement
            ratio = target_width_px / float(img.size[0])
            target_height = int(float(img.size[1]) * float(ratio))
            print(f"  Dimensions calculées: {target_width_px}x{target_height}")
            
            img.thumbnail((target_width_px, 9999), Image.Resampling.LANCZOS)
            print(f"  Après redimensionnement: {img.size}")

            # Gestion des images avec transparence (PNG)
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                print("  🔄 Conversion PNG avec transparence")
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode != 'RGB':
                print(f"  🔄 Conversion depuis {img.mode} vers RGB")
                img = img.convert('RGB')

            # Vérification finale avant sauvegarde
            if img.mode != 'RGB':
                raise ValueError(f"Mode image invalide après conversion: {img.mode}")
                
            img.save(temp_path, 'jpeg', quality=90, optimize=True)
            print(f"  ✅ Sauvegardé: {temp_path} ({os.path.getsize(temp_path)/1024:.1f} KB)")
            return temp_path

    except Exception as e:
        error_msg = f"Erreur traitement image: {type(e).__name__} - {str(e)}"
        print(f"  ❌ {error_msg}")
        log_error(img_path, error_msg)
        return None

def create_word_catalog(images_list, strate_name, output_filename):
    if not images_list:
        print(f"\n⚠️ Aucune image valide à traiter pour {strate_name}")
        return

    print(f"\n📄 Création du document Word pour {strate_name}")
    print(f"- Nombre total d'images: {len(images_list)}")
    print(f"- Nombre de produits: {len(group_images_by_product(images_list))}")
    
    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.5)

    grouped_images = group_images_by_product(images_list)

    IMAGES_PER_ROW = 2
    GUTTER_WIDTH_CM = 1.2
    page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    cell_width_cm = (page_width_cm - GUTTER_WIDTH_CM * (IMAGES_PER_ROW - 1)) / IMAGES_PER_ROW
    image_width_cm = cell_width_cm - 0.2

    # Titre principal
    title = doc.add_heading(strate_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True

    image_counter = 0
    success_count = 0
    fail_count = 0

    for product_name, product_images in grouped_images.items():
        print(f"\n📦 Produit: {product_name} ({len(product_images)} images)")
        
        # Sous-titre produit
        sub_title = doc.add_paragraph()
        run = sub_title.add_run(product_name)
        run.bold = True
        run.font.size = Pt(14)
        sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_title.paragraph_format.space_before = Pt(18)
        sub_title.paragraph_format.space_after = Pt(8)

        # Création du tableau
        num_rows = ceil(len(product_images) / IMAGES_PER_ROW)
        table = doc.add_table(rows=num_rows, cols=IMAGES_PER_ROW)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_borders_invisible(table)
        
        for col in table.columns:
            col.width = Cm(cell_width_cm)

        img_iterator = iter(product_images)
        for i in range(num_rows):
            for j in range(IMAGES_PER_ROW):
                try:
                    img_info = next(img_iterator)
                    image_counter += 1
                    
                    cell = table.cell(i, j)
                    cell.text = ''
                    cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP

                    # Légende
                    caption_text = f"{img_info['libelle_produit']} - {img_info['libelle_unite']}"
                    cap_para = cell.add_paragraph(caption_text)
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.runs[0].bold = True
                    cap_para.runs[0].font.size = Pt(10)
                    cap_para.paragraph_format.space_after = Pt(6)

                    # Traitement et insertion de l'image
                    resized_path = process_and_resize_image(
                        img_info['path'], 
                        Cm(image_width_cm).inches, 
                        image_counter
                    )

                    if resized_path and os.path.exists(resized_path):
                        para_img = cell.add_paragraph()
                        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para_img.add_run().add_picture(resized_path, width=Cm(image_width_cm))
                        success_count += 1
                        print(f"  ✅ Image #{image_counter} insérée: {os.path.basename(img_info['path'])}")
                    else:
                        fail_count += 1
                        cell.add_paragraph("[Image manquante]").alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"  ❌ Échec insertion image #{image_counter}: {os.path.basename(img_info['path'])}")

                except StopIteration:
                    pass
                except Exception as e:
                    fail_count += 1
                    error_msg = f"Erreur insertion Word: {type(e).__name__} - {str(e)}"
                    print(f"  ❌ {error_msg}")
                    log_error(img_info.get('path', 'inconnue'), error_msg)

        doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Résumé avant sauvegarde
    print(f"\n📊 Résumé final:")
    print(f"- Images traitées: {image_counter}")
    print(f"- Succès: {success_count}")
    print(f"- Échecs: {fail_count}")
    print(f"- Taux de succès: {success_count/image_counter*100:.1f}%" if image_counter > 0 else "Aucune image traitée")

    try:
        doc.save(output_filename)
        print(f"\n✅ Catalogue généré avec succès: {output_filename}")
        print(f"📏 Taille du fichier: {os.path.getsize(output_filename)/1024/1024:.2f} MB")
    except Exception as e:
        error_msg = f"Erreur sauvegarde document: {type(e).__name__} - {str(e)}"
        print(f"\n❌ {error_msg}")
        log_error(output_filename, error_msg)

if __name__ == "__main__":
    start_time = time.time()
    
    if os.path.exists(LOG_FILE):
        os.remove(LOG_FILE)
    print("🔍 Démarrage du script de génération de catalogue")

    # Préparation et nettoyage du dossier temporaire
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)
    os.makedirs(TEMP_DIR)
    print(f"📂 Dossier temporaire créé: {TEMP_DIR}")

    strates = ["101_KOLDA"]
    print("\n🚀 Démarrage du traitement des strates...")

    for strate in strates:
        strate_path = os.path.join(root_path, strate)
        if not os.path.exists(strate_path):
            print(f"⚠️ Dossier de strate introuvable: {strate_path}")
            continue
            
        images = process_strate(strate_path)
        if images:
            output_file = os.path.join(root_path, f"{strate}_catalogue_FINAL.docx")
            create_word_catalog(images, strate, output_file)
        else:
            print(f"  ❌ Aucune image valide trouvée pour la strate {strate}")

    # Nettoyage final
    try:
        shutil.rmtree(TEMP_DIR)
        print(f"\n🗑️ Dossier temporaire supprimé: {TEMP_DIR}")
    except Exception as e:
        print(f"⚠️ Impossible de supprimer le dossier temporaire: {type(e).__name__} - {str(e)}")

    # Affichage du temps d'exécution
    execution_time = time.time() - start_time
    print(f"\n⏱ Temps total d'exécution: {execution_time:.2f} secondes")
    
    print("\n✅ Génération terminée !")
    if os.path.exists(LOG_FILE) and os.path.getsize(LOG_FILE) > 0:
        print(f"📝 Journal des erreurs disponible ici: {os.path.abspath(LOG_FILE)}")
    else:
        print("ℹ️  Aucune erreur enregistrée dans le journal")