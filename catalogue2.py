import os
import glob
from PIL import Image, UnidentifiedImageError
import docx
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from math import ceil
import traceback
import shutil
import time

# === Paramètres ===
# !!! ADAPTEZ CE CHEMIN À VOTRE ORDINATEUR !!!
root_path = r"F:\projet\projet_img"
DPI = 96
LOG_FILE = "log_erreurs.txt"
MAX_IMAGE_SIZE_MB = 10
TEMP_DIR = "temp_image_processing"

def log_error(path, message):
    """Enregistre les erreurs dans un fichier log"""
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {path} => {message}\n")

def process_strate(strate_path):
    """
    Traite une strate avec structure à 3 niveaux:
    Strate -> Produit -> Unité -> Images
    """
    images = []
    if not os.path.exists(strate_path):
        print(f"❌ Dossier de strate introuvable: {strate_path}")
        return images
    
    print(f"\n▶️ Traitement de la strate : {os.path.basename(strate_path)}")
    total_images = 0
    valid_images = 0
    
    for product_dir in sorted(os.listdir(strate_path)):
        product_path = os.path.join(strate_path, product_dir)
        if not os.path.isdir(product_path):
            continue
            
        product_name = product_dir.strip()
        print(f"\n📁 Produit: {product_name}")
        
        for unit_dir in sorted(os.listdir(product_path)):
            unit_path = os.path.join(product_path, unit_dir)
            if not os.path.isdir(unit_path):
                continue
                
            unit_name = unit_dir.strip()
            print(f"  📦 Unité: {unit_name}")
            
            potential_images = []
            for file in os.listdir(unit_path):
                if file.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff')):
                    img_path = os.path.join(unit_path, file)
                    try:
                        file_size = os.path.getsize(img_path)
                        if file_size > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                            print(f"    ⚠️ Image trop grande ({file_size/1024/1024:.2f} MB): {file}")
                            log_error(img_path, f"Image trop grande ({file_size/1024/1024:.2f} MB)")
                            continue
                        potential_images.append((img_path, file_size))
                        total_images += 1
                    except OSError as e:
                        log_error(img_path, f"Erreur accès fichier: {str(e)}")
                        continue
            
            potential_images.sort(key=lambda x: x[1], reverse=True)
            found_valid_image = False
            
            for img_path, size in potential_images:
                try:
                    with Image.open(img_path) as img_test:
                        img_test.verify()
                    
                    img_test = Image.open(img_path)
                    img_test.load()
                    
                    images.append({
                        'path': img_path,
                        'libelle_groupe': product_name,
                        'libelle_produit': product_name,
                        'libelle_unite': unit_name,
                        'size': f"{img_test.size[0]}x{img_test.size[1]}",
                        'format': img_test.format
                    })
                    valid_images += 1
                    found_valid_image = True
                    print(f"    ✅ Image valide: {os.path.basename(img_path)} ({img_test.size[0]}x{img_test.size[1]})")
                    break
                    
                except Exception as e:
                    error_type = type(e).__name__
                    print(f"    ❌ Erreur image {os.path.basename(img_path)}: {error_type}")
                    log_error(img_path, f"{error_type}: {str(e)}")
                    continue
            
            if not found_valid_image and potential_images:
                print(f"    ⚠️ Aucune image valide pour {unit_name}")
                log_error(unit_path, "Aucune image valide trouvée")

    print(f"\n📊 Résumé pour {os.path.basename(strate_path)}:")
    print(f"- Images trouvées: {total_images}")
    print(f"- Images valides: {valid_images}")
    print(f"- Taux de succès: {valid_images/total_images*100:.1f}%" if total_images > 0 else "Aucune image trouvée")
    
    return images

def group_images_by_product(images_list):
    """Groupe les images par produit"""
    grouped = {}
    for img in images_list:
        product = img['libelle_produit']
        if product not in grouped:
            grouped[product] = []
        grouped[product].append(img)
    return grouped

def set_table_borders_invisible(table):
    """Supprime les bordures du tableau"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'nil')
        tblBorders.append(border_el)

def process_and_resize_image(img_path, target_width_inches, target_height_inches, unique_id, temp_dir):
    """
    Redimensionne une image, la sauvegarde dans le dossier temporaire
    et retourne son chemin ainsi que ses nouvelles dimensions en pixels.
    """
    try:
        print(f"    🔄 Traitement image #{unique_id}: {os.path.basename(img_path)}")
        temp_filename = f"temp_image_{unique_id}.jpg"
        temp_path = os.path.join(temp_dir, temp_filename)

        with Image.open(img_path) as img:
            target_width_px = int(target_width_inches * DPI)
            target_height_px= int(target_height_inches * DPI)
            
            img.thumbnail((target_width_px,target_height_px), Image.Resampling.LANCZOS)
            
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            img.save(temp_path, 'JPEG', quality=90, optimize=True)
            
            new_width_px, new_height_px = img.size
            print(f"    ✅ Image redimensionnée: {new_width_px}x{new_height_px}")
            
            return temp_path, new_width_px, new_height_px

    except Exception as e:
        error_msg = f"Erreur redimensionnement: {type(e).__name__} - {str(e)}"
        print(f"    ❌ {error_msg}")
        log_error(img_path, error_msg)
        return None, None, None

def create_word_catalog(images_list, strate_name, output_filename):
    """
    Crée le catalogue Word avec des images organisées et correctement dimensionnées.
    """
    if not images_list:
        print(f"⚠️ Aucune image à traiter pour {strate_name}")
        return

    print(f"\n📄 Création du catalogue Word pour {strate_name}")
    print(f"- Nombre d'images: {len(images_list)}")
    
    temp_dir = os.path.join(os.path.dirname(output_filename), TEMP_DIR)
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    doc = docx.Document()
    
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.5)

    title = doc.add_heading(strate_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True

    grouped_images = group_images_by_product(images_list)
    
    IMAGES_PER_ROW = 2
    page_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    cell_width_cm = page_width_cm / IMAGES_PER_ROW
    image_width_cm = cell_width_cm - 1.0
    image_height_cm= image_width_cm

    image_counter = 0
    success_count = 0
    fail_count = 0

    for product_name, product_images in grouped_images.items():
        print(f"\n📦 Produit: {product_name} ({len(product_images)} images)")
        
        sub_title = doc.add_paragraph()
        run = sub_title.add_run(product_name)
        run.bold = True
        run.font.size = Pt(14)
        sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_title.paragraph_format.space_before = Pt(18)
        sub_title.paragraph_format.space_after = Pt(8)

        num_rows = ceil(len(product_images) / IMAGES_PER_ROW)
        table = doc.add_table(rows=num_rows, cols=IMAGES_PER_ROW)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_table_borders_invisible(table)
        
        for col in table.columns:
            col.width = Cm(cell_width_cm)

        img_iterator = iter(product_images)
        for i in range(num_rows):
            for j in range(IMAGES_PER_ROW):
                try:
                    img_info = next(img_iterator)
                    image_counter += 1
                    
                    cell = table.cell(i, j)
                    cell.text = ''
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                    caption_text = f"{img_info['libelle_produit']} - {img_info['libelle_unite']}"
                    cap_para = cell.add_paragraph(caption_text)
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.runs[0].bold = True
                    cap_para.runs[0].font.size = Pt(10)
                    cap_para.paragraph_format.space_after = Pt(6)

                    # ----- DÉBUT DE LA PARTIE CORRIGÉE -----
                    resized_path, width_px, height_px = process_and_resize_image(
                        img_info['path'], 
                        Cm(image_width_cm).inches,
                        Cm(image_height_cm).inches,
                        image_counter,
                        temp_dir
                    )

                    if resized_path and os.path.exists(resized_path):
                        # Calculer la hauteur d'affichage pour conserver les proportions
                        display_width = Cm(image_width_cm)
                        # aspect_ratio = height_px / width_px
                        # display_height = display_width * aspect_ratio
                        display_height = Cm(image_height_cm)
                        
                        para_img = cell.add_paragraph()
                        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Ajouter l'image en spécifiant LA LARGEUR ET LA HAUTEUR
                        run_img = para_img.add_run()
                        run_img.add_picture(
                            resized_path, 
                            width=display_width, 
                            height=display_height
                        )
                        
                        success_count += 1
                        print(f"    ✅ Image #{image_counter} ajoutée avec dimensions forcées.")
                    # ----- FIN DE LA PARTIE CORRIGÉE -----
                    else:
                        fail_count += 1
                        cell.add_paragraph("[Image non disponible]").alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"    ❌ Échec image #{image_counter}")

                except StopIteration:
                    break
                except Exception as e:
                    fail_count += 1
                    error_msg = f"Erreur insertion: {type(e).__name__} - {str(e)}"
                    print(f"    ❌ {error_msg}")
                    log_error(img_info.get('path', 'inconnue'), error_msg)

        doc.add_paragraph("").paragraph_format.space_after = Pt(12)

    print(f"\n📊 Résumé final:")
    print(f"- Images traitées: {image_counter}")
    print(f"- Succès: {success_count}")
    print(f"- Échecs: {fail_count}")
    print(f"- Taux de succès: {success_count/image_counter*100:.1f}%" if image_counter > 0 else "Aucune image")

    try:
        doc.save(output_filename)
        print(f"\n✅ Catalogue sauvegardé: {output_filename}")
        print(f"📏 Taille: {os.path.getsize(output_filename)/1024/1024:.2f} MB")
    except Exception as e:
        error_msg = f"Erreur sauvegarde: {type(e).__name__} - {str(e)}"
        print(f"❌ {error_msg}")
        log_error(output_filename, error_msg)
    
    try:
        shutil.rmtree(temp_dir)
        print(f"🗑️ Dossier temporaire supprimé")
    except:
        pass

def main():
    """
    Fonction principale
    """
    start_time = time.time()
    
    if os.path.exists(LOG_FILE):
        os.remove(LOG_FILE)
    
    print("🚀 Démarrage de la génération de catalogue")
    print(f"📂 Répertoire racine: {root_path}")
    
    # Liste des strates à traiter (personnalisez cette liste)
    strates = [
        "101_KOLDA",
        # "102_VELINGARA",
        # "103_MEDINA YORO FOULAH",
    ]
    
    total_processed = 0
    total_errors = 0
    
    for strate in strates:
        strate_path = os.path.join(root_path, strate)
        if not os.path.exists(strate_path):
            print(f"⚠️ Dossier introuvable: {strate_path}")
            total_errors += 1
            continue
        
        try:
            images = process_strate(strate_path)
            
            if images:
                output_file = os.path.join(root_path, f"{strate}_catalogue_FINAL.docx")
                create_word_catalog(images, strate, output_file)
                total_processed += 1
            else:
                print(f"⚠️ Aucune image valide pour {strate}")
                total_errors += 1
                
        except Exception as e:
            print(f"❌ Erreur globale pour {strate}: {type(e).__name__} - {str(e)}")
            traceback.print_exc()
            log_error(strate_path, f"Erreur globale: {str(e)}")
            total_errors += 1
    
    execution_time = time.time() - start_time
    print(f"\n✅ Génération terminée!")
    print(f"📊 Statistiques:")
    print(f"   - Strates traitées: {total_processed}")
    print(f"   - Erreurs: {total_errors}")
    print(f"   - Temps d'exécution: {execution_time:.2f} secondes")
    
    if os.path.exists(LOG_FILE) and os.path.getsize(LOG_FILE) > 0:
        print(f"📝 Log des erreurs: {os.path.abspath(LOG_FILE)}")

if __name__ == "__main__":
    main()

